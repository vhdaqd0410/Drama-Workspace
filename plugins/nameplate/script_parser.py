# -*- coding: utf-8 -*-
"""
《丐世神婿》剧本解析工具
功能：导入docx剧本，自动提取前N集(默认30)首次出场的人物与场景，生成Excel表格。
用法：双击运行，或命令行 `python script_parser.py`
依赖：python-docx, openpyxl（若无，程序会自动提示/自动安装）
"""

import re
import os
import sys
import time
import subprocess

# ---------- 依赖检查 ----------
def ensure_deps():
    missing = []
    try:
        import docx  # noqa
    except ImportError:
        missing.append('python-docx')
    try:
        import openpyxl  # noqa
    except ImportError:
        missing.append('openpyxl')
    return missing


def auto_install(pkg):
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg])


# ---------- 中文数字 -> 阿拉伯数字 ----------
_ZH = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
       '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}


def zh_to_num(s):
    s = s.strip()
    if s.isdigit():
        return int(s)
    if len(s) == 1:
        return _ZH.get(s, 0)
    if len(s) == 2:
        if s[1] == '十':          # 二十 三十...
            return _ZH.get(s[0], 0) * 10
        if s[0] == '十':          # 十一 十二...
            return 10 + _ZH.get(s[1], 0)
    if len(s) == 3 and s[1] == '十':   # 二十一 三十一
        return _ZH.get(s[0], 0) * 10 + _ZH.get(s[2], 0)
    return 0


# ---------- 非主要/无名角色排除规则 ----------
# 这些是"角色类型"，不当作有名字的主要人物
NON_MAIN_PATTERNS = [
    '乞丐', '宾客', '手下', '保镖', '天策卫', '门卫', '礼仪', '司仪',
    '旗袍', '同学', '路人', '观众', '群演', '护士', '医生', '警察',
    '士兵', '仆', '侍女', '丫鬟', '家丁', '小厮', '保安', '服务员',
    '员工', '主管', '职员', '经理助手', '特使（无台词', '特使(无台词',
    '宾客（无台词', '宾客(无台词',
    '秘书', '老太君',
]

# 排除含这些字眼的整段人物行（多为"有台词/无台词"等群演标注）
GENERIC_KEYWORDS = ['（无台词', '(无台词', '（有台词', '(有台词', '有台词', '无台词']

# 文档笔误/别名统一
NAME_ALIAS = {
    '李织锦': '李知锦',   # 文档在第二十集后误写为"李织锦"
    '南屿候': '南屿侯',   # 第三十集笔误
    '山本': '山本康盛',
}


def extract_people_from_line(line, episode_num, first_episode, people_info):
    """
    从"人物："行提取有名字人物。
    line: 人物行内容（不含"人物："前缀）
    first_episode: dict {name: ep} 记录首次出场
    people_info: dict {name: {roles:[], line:首次行原文}}
    """
    # 分割：中英文逗号、顿号、空格
    tokens = re.split(r'[,，、\s]+', line.strip())
    for t in tokens:
        t = t.strip()
        if not t:
            continue
        # 跳过带数字的角色类型（如 乞丐12 / 宾客4 / 手下2）
        if re.match(r'^[一-龥]+[0-9]+$', t):
            continue
        # 跳过含"台词/无台词"标注的
        if any(k in t for k in GENERIC_KEYWORDS):
            continue
        # 排除非主要类型
        if any(p in t for p in NON_MAIN_PATTERNS):
            continue
        # 排除过长（可能是多个名字粘在一起，但保留2-4字为主的姓名；超过6字多半是描述）
        if len(t) > 6:
            continue
        # 排除纯英文/数字残留
        if re.fullmatch(r'[A-Za-z0-9]+', t):
            continue
        # 排除含标点或动作/描述词的 token（人物行不应含这些）
        if re.search(r'[。，、；：！？…【】《》"“”\s]', t):
            continue
        if t.endswith(('作揖', '行礼', '上前', '起身', '坐', '跪', '站', '退', '点头', '摇头')):
            continue
        if any(w in t for w in ('作揖', '行礼', '起身', '走进', '跪')):
            continue
        # 排除"人物"相关残留
        if t.startswith('人物'):
            continue
        # 规范化文档笔误（李织锦→李知锦 等）
        t = NAME_ALIAS.get(t, t)
        # 记录
        if t not in first_episode:
            first_episode[t] = episode_num
            people_info[t] = {'roles': [], 'line': line.strip()}
        # 若同场又出现，更新
        people_info.setdefault(t, {'roles': [], 'line': line.strip()})
        people_info[t]['roles'].append(episode_num)


# ---------- 场景提取 ----------
SCENE_RE = re.compile(r'^(日|夜|黄昏|傍晚|清晨|清晨|上午|下午)[-－](内|外|室内|室外|内景|外景)[-－]?(.+)$')


def extract_scene(line):
    m = SCENE_RE.match(line.strip())
    if m:
        loc = m.group(3).strip()
        typ = m.group(2)
        name = f"{m.group(1)}-{typ}-{loc}" if loc else f"{m.group(1)}-{typ}"
        return name, typ
    return None, None


# ---------- 身份标签 ----------
def guess_role(name, people_info, first_episode, episode_lines):
    """
    基于角色台词/上下文，给一个身份标签。
    采用启发式规则（精简版）。
    """
    roles = []
    # 台词检索
    sayings = []
    for i, ep_num in enumerate(people_info[name]['roles'][:3]):
        pass
    # 简化：用首次出场集内的台词来猜
    ep = first_episode[name]
    sayings = episode_lines.get(ep, [])
    name_talk = [l for l in sayings if l.startswith(name + '：') or l.startswith(name + ':')
                 or l.startswith(name + '（')]

    all_talk = ' '.join(sayings)

    if name == '林玄策':
        return '主角/天策神王'
    if name == '李知锦':
        return '女主角/李知锦公司总裁'
    if name in ('关白釉',):
        return '李知锦闺蜜'
    if name == '季芳非':
        return '李知锦闺蜜（反面）'
    if name in ('山本康盛',):
        return '黑金集团继承人（头号反派）'
    if name == '白泽':
        return '天策神将'
    if name == '周天成':
        return '天策神王下属/债主'
    if name == '何宏':
        return '何家继承人（反派）'
    if name == '沈盘古':
        return '盘古集团掌舵人/沈家家主'
    if name == '常经理':
        return '盘古集团经理'
    if name == '孔少泽':
        return '孔家少爷（反派）'
    if name in ('孔天成', '孔大人'):
        return '东林阁主/文官之首（反派）'
    if name in ('北境王',):
        return '四国特使·北境王'
    if name in ('南屿侯',):
        return '四国特使·南屿侯'
    if name in ('西凉特使', '西凉使'):
        return '四国特使·西凉'
    if name in ('东海公主',):
        return '四国特使·东海公主'
    if name in ('李母',):
        return '李知锦母亲'
    if name in ('李文乐',):
        return '李知锦弟弟'
    return '剧中人物'


# ---------- 精简背景 ----------
def gen_brief(name, role, first_episode, episode_lines):
    ep = first_episode[name]
    sayings = episode_lines.get(ep, [])
    # 提取首次出场集里该角色台词片段
    quotes = [l for l in sayings if re.match(r'^' + re.escape(name) + r'[：:]', l)]
    clip = ''
    if quotes:
        first_q = quotes[0]
        # 去掉"名字："前缀，取前半句
        content = re.sub(r'^' + re.escape(name) + r'[：:]', '', first_q)
        clip = content[:28]
    elif any(re.match(r'^' + re.escape(name), l) for l in sayings):
        first_q = [l for l in sayings if re.match(r'^' + re.escape(name), l)][0]
        content = re.sub(r'^' + re.escape(name), '', first_q)
        clip = content[:28]

    if clip:
        return f'{role}。第{zh_num_to_cn(ep)}集登场，首现台词：“{clip}…”'
    return f'{role}。第{zh_num_to_cn(ep)}集登场。'


def zh_num_to_cn(n):
    if n == 0:
        return '?'
    digits = '一二三四五六七八九'
    if n <= 10:
        return digits[n-1] if n < 10 else '十'
    if n < 20:
        return '十' + digits[n-11] if n > 10 else '十'
    if n < 100:
        tens = n // 10
        ones = n % 10
        return (digits[tens-1] + '十') + (digits[ones-1] if ones else '')
    return str(n)


# ---------- 准确率校准 ----------
# 配置文件保存在工具目录下 calibration.json
CALIB_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'calibration.json')


def _resource_path(name):
    """获取资源文件路径（兼容 PyInstaller 打包后的 exe）。"""
    import sys
    base = getattr(sys, '_MEIPASS', None)
    if base:
        return os.path.join(base, name)
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), name)


def default_calibration():
    return {'exclude': [], 'aliases': {}, 'roles': {}}


def load_calibration():
    """读取校准配置。若文件不存在返回默认空配置。"""
    try:
        if os.path.exists(CALIB_FILE):
            import json
            with open(CALIB_FILE, 'r', encoding='utf-8') as f:
                cfg = json.load(f)
            cfg.setdefault('exclude', [])
            cfg.setdefault('aliases', {})
            cfg.setdefault('roles', {})
            return cfg
    except Exception:
        pass
    return default_calibration()


def save_calibration(cfg):
    """保存校准配置到文件。"""
    import json
    try:
        with open(CALIB_FILE, 'w', encoding='utf-8') as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def _norm_name(s):
    return (s or '').strip().lower()


def apply_calibration(data, cfg=None):
    """把校准规则应用到解析结果 data。
    - exclude: 排除这些名字（中文名或英文名匹配）
    - aliases: {旧名: 规范名}，把别名统一/合并（如 George -> 乔治）
    - roles:   {名字: 身份}，覆盖角色身份
    直接修改并返回 data。"""
    if cfg is None:
        cfg = load_calibration()
    if not data:
        return data
    exclude = set(_norm_name(x) for x in cfg.get('exclude', []))
    aliases = cfg.get('aliases', {}) or {}
    roles = cfg.get('roles', {}) or {}

    people = data.get('people') or []
    kept = []
    for item in people:
        # 解包（兼容4/5元素）
        if len(item) == 5:
            cn, en, ep, role, brief = item
        else:
            cn, ep, role, brief = item
            en = ''
        # 排除
        if cn and _norm_name(cn) in exclude:
            continue
        if en and _norm_name(en) in exclude:
            continue
        # 别名统一（中文名和英文名都尝试）
        new_cn = cn
        new_en = en
        if cn and _norm_name(cn) in aliases:
            new_cn = aliases[_norm_name(cn)]
        elif en and _norm_name(en) in aliases:
            new_cn = aliases[_norm_name(en)]
        # 角色身份覆盖
        if cn and _norm_name(cn) in roles:
            role = roles[_norm_name(cn)]
        elif en and _norm_name(en) in roles:
            role = roles[_norm_name(en)]
        if len(item) == 5:
            kept.append([new_cn, new_en, ep, role, brief])
        else:
            kept.append([new_cn, ep, role, brief])
    data['people'] = kept
    return data


# ---------- 主解析 ----------
def parse_script(docx_path, max_episodes=30, lang_mode='auto'):
    import docx

    d = docx.Document(docx_path)
    lines = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    # 也考虑表格文本
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(t)

    # 检测剧本类型
    text_all = '\n'.join(lines)

    # 英文正文剧本：场景 "1-1 日 内 地点" 或 "1-1 地点 / 日 / 内"，且含英文台词 "Name（对X）："
    has_scene_digit = re.search(r'^\d+-\d+\s+', text_all, re.M)
    has_en_talk = re.search(r'^[A-Za-z][A-Za-z .\'\-]*\s*[（(][^）)]*[）)]\s*[：:]\s*\w', text_all, re.M)
    has_en_char = re.search(r'^人物[：:]\s*[A-Za-z]', text_all, re.M)

    # 类型B：人物小传式（含"出场集数"）
    is_duallang = '人物小传' in text_all and '出场集数' in text_all
    # 英文正文剧本：有数字场景标题 + 英文台词 + (人物行用英文名 或 台词明显英文)
    is_english_body = bool(has_scene_digit and has_en_talk)

    if is_duallang:
        data, err = parse_script_duallang(lines, max_episodes)
        if data:
            data['title'] = d.core_properties.title or os.path.basename(docx_path)
        return data, err
    if lang_mode in ('en',) or (is_english_body and lang_mode != 'zh'):
        data, err = parse_script_english(lines, max_episodes)
        if data:
            data['title'] = d.core_properties.title or os.path.basename(docx_path)
        return data, err

    # ---------- 以下为类型A（中文剧本） ----------
    # 定位集数
    ep_starts = []  # (num, line_index)
    for idx, l in enumerate(lines):
        m = re.match(r'^第([一二三四五六七八九十百\d]+)集', l)
        if m:
            ep_starts.append((zh_to_num(m.group(1)), idx))
    ep_starts.sort()
    if not ep_starts:
        return None, '未找到集数标题，请确认剧本格式'

    # 找到第 max_episodes 集的结束
    end_idx = len(lines)
    for num, idx in ep_starts:
        if num == max_episodes:
            # 该集结束于下一集
            for num2, idx2 in ep_starts:
                if num2 > max_episodes:
                    end_idx = idx2
                    break
            break

    # 计算每集的行范围 [start, end)
    ep_ranges = {}
    for num, idx in ep_starts:
        if num > max_episodes:
            break
        nxt = end_idx
        for num2, idx2 in ep_starts:
            if num2 == num + 1:
                nxt = idx2
                break
        ep_ranges[num] = (idx, nxt)

    # 组织每集的文本行
    episode_lines = {}   # {ep: [lines]}
    episode_text = {}    # {ep: text}
    for num, (idx, nxt) in ep_ranges.items():
        episode_lines[num] = lines[idx:nxt]
        episode_text[num] = '\n'.join(lines[idx:nxt])

    # ---------- 人物收集（综合多种来源） ----------
    # 方法A：收集所有"人物："行出现的人物名字，记首次集数
    # 方法B：收集正文中 "名字：台词" 和正文名字出现，补全没被"人物："行收录的重要角色
    first_episode = {}      # name -> 首次集数
    people_info = {}        # name -> {roles:[], line:首次人物行, appears:[]}

    def record_people_from_line(content, ep):
        extract_people_from_line(content, ep, first_episode, people_info)

    # 方法A：人物行
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            l = lines[i]
            if re.match(r'^人物[：:]', l):
                content = re.sub(r'^人物[：:]', '', l)
                record_people_from_line(content, num)

    # 方法B：正文台词 "名字：..."，捕捉没进"人物："行的角色（如沈盘古）
    # 噪声词：脚本标注类、非人物
    NOISE_SUFFIX = ['VO', 'OV', 'OS', 'ov', 'os', 'vo', 'OV', 'ov']
    NOISE_WORDS = ['特效', '特写', '镜头', '众人', '旁白', '音效', '慢镜头',
                   '俯拍', '仰拍', '画外音', '字幕', 'BGM', '音效', '背景',
                   '人物', '场景', '灯光', '道具', '服装', '导演', '摄影',
                   '林母', '统领', '司仪', '礼仪小姐', '门卫', '配音', '演员']
    # 已知主要人物（来自方法A）
    known_main = set(first_episode.keys())
    # 动作/状态后缀（"白泽惊"= 白泽+惊）
    ACTION_SUFFIX = ['惊', '怒', '笑', '问', '道', '冷', '急', '慌', '喜',
                     '叹', '愣', '皱眉', '沉', '讽', '讥', '吼', '骂', '惧',
                     '淡然', '冷笑', '嗤笑', '大笑', '惊呼', '震怒', '错愕',
                     '跳', '起身', '回', '看向', '指', '瞪']
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            l = lines[i]
            m = re.match(r'^([一-龥A-Za-z]{1,6})[：:]\s*(.+)$', l)
            if not m:
                continue
            name = m.group(1).strip()
            # 过滤：非主要类型
            if any(p in name for p in NON_MAIN_PATTERNS):
                continue
            if re.search(r'[0-9]', name):
                continue
            if any(w in name for w in NOISE_WORDS):
                continue
            if any(name.endswith(s) for s in NOISE_SUFFIX):
                continue
            if name in ('VO', 'OV', 'OS', 'os', 'vo', 'ov'):
                continue
            if len(name) > 6:
                continue
            # 排除含动作/标点/描述的 name（如"白泽作揖行礼"）
            if re.search(r'[。，、；：！？…【】]', name):
                continue
            if any(w in name for w in ('作揖', '行礼', '起身', '走进', '跪下', '上前', '伸出', '抱拳', '退下', '低头')):
                continue
            if name.endswith(('作揖', '行礼', '起身', '跪下', '上前')):
                continue
            # 已知人物+动作后缀（如"白泽惊"），视为该人物的动作标注
            if name not in known_main:
                is_action = False
                for km in known_main:
                    if len(km) >= 2 and name.startswith(km) and name != km:
                        tail = name[len(km):]
                        if any(a == tail or tail.startswith(a) for a in ACTION_SUFFIX):
                            is_action = True
                            break
                if is_action:
                    continue
            # 规范化文档笔误：李织锦==李知锦，南屿候==南屿侯
            name = NAME_ALIAS.get(name, name)
            # 记录并允许用更早集数更新（如常经理第9集已有台词，但方法A在第10集人物行才记录）
            if name not in first_episode or num < first_episode[name]:
                first_episode[name] = num
                people_info.setdefault(name, {'roles': [], 'line': '', 'appears': []})

    # 方法C：识别"正式登场"（画面中实际出现），排除仅在别人台词中被提及
    # 正式登场信号：
    #   1) 该名字独立说台词 "名字："（已在方法B处理）
    #   2) 该名字出现在 "△..." 场景描述中作为登场主体（如"△白泽坐在车内"、"△沈盘古带领手下走进来"）
    # 为避免把"被提及"当成"登场"，△行内若名字前有 看到/提到/望向/谈起/说起/想起/传来/提及 等引出词，则不视为登场。
    TRIGGER_PREFIX = ['看到', '望向', '提到', '说起', '谈起', '想起', '提及',
                      '传来', '听说', '知道', '认为', '想起', '看向', '只见',
                      '抬眼', '望向', '看着', '望见', '看着']
    # 先收集所有名字（当前first_episode已含方法A+B的人物）
    all_names = list(first_episode.keys())
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            l = lines[i]
            if not l.startswith('△'):
                continue
            for name in all_names:
                if name not in l:
                    continue
                # 名字出现位置前的一段，判断是否被"引出词"修饰（视为被提及而非登场）
                pos = l.find(name)
                if pos == -1:
                    continue
                before = l[:pos]
                # 名字前若有主语引出词且该词紧跟名字，多半是被提及
                triggered = any(before.rstrip().endswith(tp) for tp in TRIGGER_PREFIX)
                if triggered:
                    continue
                # 名字前是"的/了/向/对/与/和"等，且非独立主体——保守起见：若名字前一个字是"的/向/对/被/把/让/给"则视为被提及
                if before and before[-1] in '的向对被把让给将叫从望':
                    continue
                # 该名字在该集正式登场
                if num < first_episode[name]:
                    first_episode[name] = num

    # 提取场景（首次出现）
    scene_first = {}
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            name, typ = extract_scene(lines[i])
            if name and name not in scene_first:
                scene_first[name] = (num, typ)

    # 生成人物最终数据
    people_final = []
    for name, ep in sorted(first_episode.items(), key=lambda x: x[1]):
        role = guess_role(name, people_info, first_episode, episode_lines)
        brief = gen_brief(name, role, first_episode, episode_lines)
        people_final.append((name, ep, role, brief))

    # 生成场景最终数据
    scene_final = []
    for name, (ep, typ) in sorted(scene_first.items(), key=lambda x: x[1][0]):
        scene_final.append((name, ep, typ))

    # 提取字幕
    subtitles = _extract_subtitles(lines, ep_ranges, max_episodes)

    return {
        'title': d.core_properties.title or os.path.basename(docx_path),
        'people': people_final,
        'scenes': scene_final,
        'subtitles': subtitles,
    }, None


def _extract_subtitles(lines, ep_ranges, max_episodes):
    """统一提取【字幕：...】中属于'时间注释'的字幕，返回 [(集数, 内容)]。"""
    if ep_ranges is None:
        return []
    subs = []
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            l = lines[i]
            m = re.search(r'【\s*字幕[：:]?\s*([^】]+)】', l)
            if m:
                content = m.group(1).strip()
                if content and _is_time_subtitle(content):
                    subs.append((num, content))
    return subs


def _is_time_subtitle(content):
    """判断一条字幕是否属于'需要添加到视频上的时间注释'（而非角色介绍/道具说明等）。"""
    import re
    low = content.lower().strip()
    if not low:
        return False
    # 明确的时间副词/时间指示短语 → 时间类
    time_marker = [
        ' ago', ' later', 'next day', 'next night', 'the next day', 'the next night',
        'that day', 'that night', 'the same day', 'the same night', 'that morning',
        'this morning', 'that afternoon', 'this afternoon', 'that evening', 'this evening',
        'tonight', 'today', 'tomorrow', 'yesterday', 'meanwhile', 'at the same time',
        'earlier', 'after a while', 'a moment later', 'the following day',
        'a few days', 'a few hours', 'a few minutes', 'a few moments', 'a few weeks',
        'a few months', 'a few years', 'years later', 'years ago', 'days later', 'days ago',
        'hours later', 'minutes later', 'weeks later', 'months later', 'weeks ago',
        'months ago', 'hours ago', 'minutes ago', 'century', 'decade', 'decades',
        'one week', 'two weeks', 'one month', 'two months', 'one year', 'two years',
        'one day', 'two days', 'one hour', 'two hours', 'one minute', 'two minutes',
    ]
    if any(t in low for t in time_marker):
        return True
    # 星期/时间点 → 时间类（纯时间，不含角色名）
    weekday = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
    if any(w == low.strip() or low.startswith(w + ',') or low.endswith(w) or low.startswith(w + ' ') for w in weekday):
        return True
    # 纯时间点（如 "Day 3" "Night" "3:00" "Noon" "Midnight"）
    if re.fullmatch(r'\d+[:：]\d+', low) or low in ('day', 'night', 'morning', 'afternoon', 'evening', 'noon', 'midnight', 'dawn', 'dusk', 'day 1', 'day 2'):
        return True
    # 含数字 + 时间单位，且不含"人名描述"特征（逗号后是身份描述则排除）
    has_time_unit = re.search(r'\d+\s*(years?|days?|hours?|minutes?|weeks?|months?|decades?|centur)', low)
    if has_time_unit:
        # 年龄介绍（"Luna, 5 years old"）不是时间注释，排除
        if re.search(r'years?\s+old|year\s+old', low):
            return False
        # 若含角色身份词则视为角色介绍，排除
        if re.search(r'master|heir|mage|academy|dean|chief|appraiser|instructor|princess|queen|elf|marquis|elite|student|head of|mistress|duke|lord|lady|son of|daughter of', low):
            return False
        return True
    # 其他情况视为非时间注释（角色介绍/道具说明等）
    return False


# ---------- 类型B：中英双语/人物小传式剧本解析 ----------
# 适用于含"人物小传"+"出场集数"字段的剧本
# 提取范围：人物小传主要人物 + 正文中真正开口有台词的角色（排除无名群演）
# 群演/职业名词（不是有名字的人物，正文说话但不作为角色条目）
DUALLANG_GENERIC_EN = [
    'Adjudicator', 'Auctioneer', 'Auction Manager', 'Auction Envoy',
    'Auction Clerk', 'Auction Director', 'Butler', 'Guard', 'Academy Guard Captain',
    'Captain', 'Officer', 'Clerk', 'Manager', 'Assistant', 'Staff', 'Messenger',
    'Envoy', 'Herald', 'Servant', 'Maid', 'Attendant', 'Crowd', 'VO', 'OV', 'OS',
    'Narrator', 'Guard Captain', 'Chairman', 'Secretary', 'Doctor', 'Nurse',
]
# 正文台词中短名→全名（用于把正文的 first-name 归并到小传全名）
def _normalize_en(s):
    return re.sub(r'\s+', ' ', s.strip())

# 常见配角英文名 → 中文名（音译映射，用于补充配角中文名）
DUALLANG_EXTRA_CN = {
    'gideon': '吉迪恩', 'kevin': '凯文', 'morris': '莫里斯', 'vera': '薇拉',
    'simon': '西蒙', 'nixia': '妮克莎', 'marlo': '马洛', 'grayson': '格雷森',
    'wendell': '温德尔', 'howard': '霍华德', 'damian': '达米安',
    'adrian': '亚德里安', 'victor': '维克多', 'edrick': '埃德里克',
    'serafina': '塞拉菲娜', 'luna': '露娜', 'ryan': '莱恩', 'irene': '艾琳',
    'herbert': '赫伯特', 'augustus': '奥古斯',
    'owen': '欧文', 'alice': '爱丽丝', 'ian': '伊恩', 'jason': '杰森',
    'eva': '伊娃', 'lily': '莉莉', 'bea': '贝雅', 'masson': '梅森',
    'maeve': '梅芙', 'meave': '梅芙', 'maive': '梅芙', 'stargazer': '星见者',
    'george': '乔治', 'nina': '妮娜',
}

def parse_script_duallang(lines, max_episodes=30):
    import re

    # 1) 解析人物小传：主要人物（中文名 + 英文全名 + 首次集数 + 背景）
    cast = []  # [{cn, en, first_ep, img}]
    for i, l in enumerate(lines):
        m = re.match(r'^\d+\s*[，,、]\s*(.+?)\s*[（(](.+?)[）)]\s*$', l)
        if not m:
            continue
        cn = m.group(1).strip()
        en = m.group(2).strip()
        if not re.search(r'[A-Za-z]', en):
            continue
        ep_field = ''
        img_field = ''
        for j in range(i + 1, min(i + 8, len(lines))):
            if lines[j].startswith('出场集数') and not ep_field:
                ep_field = lines[j]
            if lines[j].startswith('人物形象') and not img_field:
                img_field = lines[j]
        if not ep_field:
            continue
        nums = [int(x) for x in re.findall(r'(\d+)-', ep_field)]
        first_ep = min(nums) if nums else 999
        cast.append({
            'cn': cn, 'en': en, 'first_ep': first_ep,
            'img': img_field[5:].strip() if img_field else '',
        })

    # 建立 英文名→小传人物 映射（同时支持全名和首词）
    en2cast = {}          # 规范化英文名(小写) -> cast下标
    en_first2cast = {}    # 英文首词(小写) -> cast下标
    for ci, c in enumerate(cast):
        en_n = _normalize_en(c['en']).lower()
        en2cast[en_n] = ci
        if ' ' in en_n:
            en_first2cast[en_n.split()[0]] = ci

    # 2) 扫描正文，提取所有"英文名（括注）：台词"的说话角色及其首次集数
    #    先建立 集号 -> 行范围
    ep_starts = []
    for i, l in enumerate(lines):
        m = re.match(r'^第([一二三四五六七八九十\d]+)集', l)
        if m:
            ep_starts.append((zh_to_num(m.group(1)), i))
    ep_starts.sort()
    end30 = len(lines)
    for num, idx in ep_starts:
        if num == max_episodes:
            for num2, idx2 in ep_starts:
                if num2 > max_episodes:
                    end30 = idx2
                    break
            break
    # 当前集号（随行推进）
    cur_ep = 0
    body_speakers = {}  # 英文名 -> {first_ep, count, first_line}
    for i, l in enumerate(lines):
        if l.startswith('第') and re.match(r'^第([一二三四五六七八九十\d]+)集', l):
            cur_ep = zh_to_num(re.match(r'^第([一二三四五六七八九十\d]+)集', l).group(1))
            continue
        if cur_ep == 0 or cur_ep > max_episodes:
            continue
        m = re.match(r'^([A-Za-z][A-Za-z .\'\-]{1,30})[（(][^）)]*[）)]\s*[：:]\s*', l)
        if m:
            en = _normalize_en(m.group(1))
            if en.lower() in DUALLANG_GENERIC_EN or en.lower() in [g.lower() for g in DUALLANG_GENERIC_EN]:
                continue
            if en not in body_speakers:
                body_speakers[en] = {'first_ep': cur_ep, 'count': 0}
            body_speakers[en]['count'] += 1

    # 3) 汇总人物
    people_final = []
    used_cn = set()
    # 3a) 主要人物（小传），保留首集<=max的
    for c in cast:
        if c['first_ep'] > max_episodes:
            continue
        role = guess_role_duallang(c['cn'], c['en'])
        img = c['img']
        if img:
            brief = f"{role}。第{zh_num_to_cn(c['first_ep'])}集登场。{img[:60]}{'…' if len(img)>60 else ''}"
        else:
            brief = f"{role}。第{zh_num_to_cn(c['first_ep'])}集登场。"
        people_final.append([c['cn'], c['en'], c['first_ep'], role, brief])
        used_cn.add(c['cn'])

    # 3b) 正文有台词、但不在小传主要人物里的角色（有名有姓的配角）
    #     尝试把英文名映射回中文名：若英文名(或首词)在小传全名中→已并入；否则用音译映射
    en2cn_extra = _build_extra_en2cn(lines)
    added = {}
    for en, info in body_speakers.items():
        en_l = en.lower()
        # 跳过已归属小传主要人物
        if en_l in en2cast or en_l.split()[0] in en_first2cast:
            continue
        if info['first_ep'] > max_episodes:
            continue
        # 排除无名形式："某人的儿子/女儿"、"Grayson's Son" 等
        if re.search(r"['’]s\s+(Son|Daughter|Son-in-law|Wife|Husband|Brother|Sister|Father|Mother|Mr\.|Mrs\.)", en, re.I):
            continue
        if len(en.split()) > 2 and 'Son' in en.split()[-1]:
            continue
        # 中文名：优先"其他配角"映射，其次内置音译表，否则留空(仅显示英文名)
        cn = en2cn_extra.get(en_l) or en2cn_extra.get(en_l.split()[0]) or DUALLANG_EXTRA_CN.get(en_l.split()[0]) or ''
        if cn and cn in used_cn:
            continue
        disp_cn = cn if cn else ''
        disp_en = en
        key = disp_cn or disp_en
        if key in added:
            continue
        added[key] = True
        role = '配角'
        brief = f"配角。第{zh_num_to_cn(info['first_ep'])}集登场，正文中有台词。"
        people_final.append([disp_cn, disp_en, info['first_ep'], role, brief])

    # 按首集排序
    people_final.sort(key=lambda x: x[2])

    # 4) 提取场景：格式 "1-1 场景名 / 日 / 内"
    scene_first = {}
    for i, l in enumerate(lines):
        m = re.match(r'^(\d+)-(\d+)\s+(.+?)\s*/\s*(日|夜|清晨|黄昏|傍晚|上午|下午)\s*/\s*(内|外)', l)
        if m:
            ep = int(m.group(1))
            if ep > max_episodes:
                continue
            loc = m.group(3).strip()
            typ = m.group(5)
            if loc not in scene_first:
                scene_first[loc] = (ep, typ)

    scene_final = []
    for name, (ep, typ) in sorted(scene_first.items(), key=lambda x: x[1][0]):
        scene_final.append((name, ep, typ))

    # 提取字幕（按集号定位）
    subtitles = []
    cur_ep = 0
    for i, l in enumerate(lines):
        if re.match(r'^第([一二三四五六七八九十\d]+)集', l):
            cur_ep = zh_to_num(re.match(r'^第([一二三四五六七八九十\d]+)集', l).group(1))
            continue
        m = re.search(r'【\s*字幕[：:]?\s*([^】]+)】', l)
        if m and cur_ep and cur_ep <= max_episodes:
            c = m.group(1).strip()
            if c and _is_time_subtitle(c):
                subtitles.append((cur_ep, c))

    return {
        'title': '',
        'people': [tuple(p) for p in people_final],
        'scenes': scene_final,
        'subtitles': subtitles,
    }, None


def _build_extra_en2cn(lines):
    """从'三，其他配角'及正文里，尽量把英文名映射回中文名。返回 {英文名小写: 中文名}"""
    # 主要手段：正文中出现过 "中文名（英文名）" 的括注形式
    mapping = {}
    for l in lines:
        # 形如 露娜（Luna） 莱恩（Ryan）
        for m in re.finditer(r'([一-鿿·]{2,8})\s*[（(]([A-Za-z][A-Za-z .\'\-]*)[）)]', l):
            cn = m.group(1)
            en = _normalize_en(m.group(2))
            mapping[en.lower()] = cn
            mapping[en.lower().split()[0]] = cn
    return mapping


# 英文剧本中的群演/非主要角色词（排除）
EN_GENERIC = {
    '精灵若干', '侍从若干', '守卫若干', '宾客若干', '群众若干', '精灵们', '士兵若干',
    '若干精灵', '若干守卫', '若干侍从', '若干宾客', '围观群众', '酒馆客人', '路人',
    'Al.', 'Some elves', 'Several elves', 'Guards', 'Servants', 'Crowd', 'Elves',
    # 英文职业/群演词（含大小写变体，过滤时统一转小写匹配）
    'researchers', 'researcher', 'researcher a', 'researcher b', 'research',
    'bodyguard', 'bodyguards', 'guard', 'guards', 'servant', 'servants',
    'butler', 'maid', 'maids', 'guests', 'guest', 'journalists', 'journalist',
    'reporter', 'reporters', 'assistant', 'assistants', 'staff', 'clerks',
    'waitress', 'waiter', 'drivers', 'driver', 'taxi driver', 'white',
    'black', 'police', 'officer', 'captain', 'nurse', 'doctor', 'children',
    'crowd', 'people', 'elder', 'manager', 'secretary', 'attendant', 'messenger',
    'characters', 'all', 'everyone', 'men', 'women', 'kids', 'voices',
    'several', 'classmates', 'woman', 'party', 'officers', 'nobility',
    'computer', 'taxi', 'nobody', 'people', 'child', 'family', 'group',
    'company', 'team', 'staff', 'researcher1', 'researcher2', 'assistant1',
}

# 英文角色名中需过滤的后缀（OS/VO/OV 等）
EN_SUFFIX = (' os', ' vo', ' ov', ' os ', ' vo ', ' ov ', ' os)', ' vo)', ' ov)')


# ---------- 类型C：英文正文剧本解析 ----------
# 场景格式 "1-1 日 内 地点" 或 "1-1. INT. Lab - DAY"，人物行 "人物：Owen" 或 "Characters: SERENA"
# 台词 "Owen（对Ian）：..." 或 "Serena（落寞）OS：..."
# 输出以英文角色名为主，并提取【字幕】内容
def parse_script_english(lines, max_episodes=30):
    import re

    # 建立 集号 -> 行范围
    ep_starts = []
    for i, l in enumerate(lines):
        m = re.match(r'^第([一二三四五六七八九十\d]+)集', l)
        if m:
            ep_starts.append((zh_to_num(m.group(1)), i))
    ep_starts.sort()
    if not ep_starts:
        return None, '未找到集数标题'
    end_idx = len(lines)
    for num, idx in ep_starts:
        if num == max_episodes:
            for num2, idx2 in ep_starts:
                if num2 > max_episodes:
                    end_idx = idx2
                    break
            break
    ep_ranges = {}
    for num, idx in ep_starts:
        if num > max_episodes:
            break
        nxt = end_idx
        for num2, idx2 in ep_starts:
            if num2 == num + 1:
                nxt = idx2
                break
        ep_ranges[num] = (idx, nxt)

    # 0) 解析人物小传（若有）：中文名（英文名）：... → 建立英文名→中文名映射
    en2cn = {}   # 英文名小写 -> 中文名
    cn2en = {}   # 中文名 -> 英文名(规范)
    for l in lines:
        m = re.match(r'^([一-鿿·]{2,10})\s*[（(]\s*([A-Za-z][A-Za-z .\'\-]*)\s*[）)]\s*[：:]\s*', l)
        if m:
            cn = m.group(1).strip()
            en = _normalize_en(m.group(2))
            en2cn[en.lower()] = cn
            en2cn[en.lower().split()[0]] = cn
            cn2en[cn] = en

    first_episode = {}   # key -> 首集
    name_disp = {}       # key -> 规范显示名（英文或"cn:中文"）
    # 1) 从"人物："行或"Characters:"行提取角色
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            l = lines[i]
            is_char_line = re.match(r'^(人物|Characters|CHARACTERS)[：:]', l)
            if not is_char_line:
                continue
            content = re.sub(r'^(人物|Characters|CHARACTERS)[：:]', '', l)
            # 先移除所有括注段 (...)，避免括注内的描述词被误当角色
            content = re.sub(r'[（(][^）)]*[）)]', '', content)
            # 逐 token 处理
            for tok in re.split(r'[,，、\s]+', content):
                tok = tok.strip()
                if not tok:
                    continue
                # 过滤属格 'S / 's
                if re.search(r"'[Ss]$", tok):
                    continue
                tlow = tok.lower()
                if tlow in EN_GENERIC:
                    continue
                if re.search(r'若干|守卫|侍从|宾客|群众|围观|路人|研究员|助理|司机|侍者|女仆|管家|同学|女|男人|女人|官员|贵族', tok):
                    continue
                if re.search(r'\b(father|mother|wife|husband|son|daughter|brother|sister|child|kids|men|women)\b', tlow):
                    continue
                if re.fullmatch(r'[a-z]', tlow):
                    continue  # 单字母词（to/a）
                if re.search(r'[A-Za-z]', tok):
                    key = tlow
                    clean = re.sub(r'\s*(OS|VO|OV)$', '', tok, flags=re.I).strip()
                    if clean != tok and clean:
                        tok = clean
                        key = clean.lower()
                    if key in EN_GENERIC or key.endswith(('os', 'vo', 'ov')):
                        continue
                    if key not in first_episode:
                        first_episode[key] = num
                        name_disp[key] = tok
                elif re.search(r'[一-鿿]', tok):
                    if re.search(r'若干|守卫|侍从|宾客|群众|路人', tok):
                        continue
                    key = 'cn:' + tok
                    if key not in first_episode:
                        first_episode[key] = num
                        name_disp[key] = tok

    # 2) 从台词提取说话人，补全没进人物行的角色
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            l = lines[i]
            m = re.match(r'^([A-Za-z][A-Za-z .\'\-]{0,30})\s*[（(][^）)]*[）)]\s*[：:]\s*', l)
            if not m:
                m = re.match(r'^([A-Za-z][A-Za-z .\'\-]{0,30})[：:]\s*', l)
            if not m:
                continue
            name = m.group(1).strip()
            tlow = name.lower()
            if tlow in EN_GENERIC:
                continue
            if re.search(r'\b(father|mother|wife|husband|son|daughter|brother|sister|driver|servant|butler|maid|guest|journalist|assistant)\b', tlow):
                continue
            # 清理 OS/VO/OV 后缀
            clean = re.sub(r'\s*(OS|VO|OV)$', '', name, flags=re.I).strip()
            key = clean.lower() if clean else tlow
            if key in EN_GENERIC or key.endswith(('os', 'vo', 'ov')):
                continue
            if key not in first_episode:
                first_episode[key] = num
                name_disp[key] = clean or name

    # 3) 排除噪声
    NOISE_EN = {'vo', 'ov', 'os', 'bgm', 'sfx', 'vfx', 'narrator', 'screen',
                'shot', 'cut', 'lens', 'close-up', 'camera', 'scene', 'title',
                'subtitle', 'flashback', 'preview', 'episode', 'voice', 'characters',
                'research', 'researchers', 'serena os', 'elias os', 'serena vo',
                'elias vo', 'julian vo', 'julian os', 'mark vo', 'cordelia vo',
                'cordelia os', 'eloise os'}
    for key in list(first_episode.keys()):
        disp = name_disp[key]
        if re.search(r'\d', disp):
            del first_episode[key]
            continue
        if key in NOISE_EN:
            del first_episode[key]
            continue
        if len(disp) > 25:
            del first_episode[key]
            continue

    # 4) 提取场景：中文格式 + 英文 INT/EXT 格式
    scene_first = {}
    for i, l in enumerate(lines):
        typ = ''
        ep = 0
        loc = ''
        # 中文格式A: "1-1 日 内 地点"
        m = re.match(r'^(\d+)-(\d+)\s+(日|夜|清晨|黄昏|傍晚|上午|下午|白日)\s+(内|外)\s+(.+)$', l)
        if m:
            ep = int(m.group(1))
            typ = m.group(4)
            loc = m.group(5).strip()
        # 中文格式B: "1-1 地点 / 日 / 内"
        if not m:
            m = re.match(r'^(\d+)-(\d+)\s+(.+?)\s*/\s*(日|夜|清晨|黄昏|傍晚|上午|下午)\s*/\s*(内|外)', l)
            if m:
                ep = int(m.group(1))
                loc = m.group(3).strip()
                typ = m.group(5)
        # 英文格式: "1-1. INT. Lab Corridor - DAY"
        if not m:
            m = re.match(r'^(\d+)-(\d+)[.\s]+(INT|EXT)[.\s]+(.+?)\s*-\s*(DAY|NIGHT|DUSK|DAWN|MORNING|AFTERNOON|EVENING|NOON|LATE NIGHT|EARLY MORNING)\s*$', l)
            if m:
                ep = int(m.group(1))
                inout = m.group(3).upper()
                loc = m.group(4).strip()
                typ = '内' if inout == 'INT' else '外'
        if not m or not loc:
            continue
        if ep > max_episodes:
            continue
        if loc not in scene_first:
            scene_first[loc] = (ep, typ)

    # 5) 提取字幕：【字幕：...】（仅保留时间注释类）
    subtitles = []   # [(集数, 字幕内容)]
    for num, (idx, nxt) in ep_ranges.items():
        for i in range(idx, nxt):
            l = lines[i]
            m = re.search(r'【\s*字幕[：:]?\s*([^】]+)】', l)
            if m:
                content = m.group(1).strip()
                if content and _is_time_subtitle(content):
                    subtitles.append((num, content))

    # 6) 汇总人物
    people_final = []
    for key in sorted(first_episode.keys(), key=lambda k: first_episode[k]):
        ep = first_episode[key]
        if key.startswith('cn:'):
            cn = name_disp[key]
            en = cn2en.get(cn, '')
            role = '剧中人物'
            brief = f"第{zh_num_to_cn(ep)}集登场。"
        else:
            en = name_disp[key]
            cn = en2cn.get(key, '') or DUALLANG_EXTRA_CN.get(key.split()[0]) or ''
            role = '剧中人物'
            brief = f"第{zh_num_to_cn(ep)}集登场。"
        people_final.append([cn, en, ep, role, brief])

    scene_final = []
    for name, (ep, typ) in sorted(scene_first.items(), key=lambda x: x[1][0]):
        scene_final.append((name, ep, typ))

    return {'title': '', 'people': people_final, 'scenes': scene_final,
            'subtitles': subtitles}, None

    scene_final = []
    for name, (ep, typ) in sorted(scene_first.items(), key=lambda x: x[1][0]):
        scene_final.append((name, ep, typ))

    return {'title': '', 'people': people_final, 'scenes': scene_final}, None


def guess_role_duallang(cn, en):
    """给双语剧本人物一个身份定位（基于常见设定启发式）"""
    if cn == '露娜' or 'Luna' in en:
        return '女主角/五岁魔导师'
    if cn.startswith('莱恩') or 'Ryan' in en:
        return '男主角/兰斯特家族继承人'
    if cn == '艾琳' or 'Irene' in en:
        return '兰斯特家族核心成员'
    if cn == '赫伯特' or 'Herbert' in en:
        return '兰斯特家族老公爵/祖父'
    if cn == '埃德里克' or 'Edrick' in en:
        return '皇家魔法学院院长'
    if cn == '奥古斯' or 'Augustus' in en:
        return '天空魔法塔之主/师父'
    if cn.startswith('塞拉菲娜') or 'Serafina' in en:
        return '布莱克伍德家继承人/学院首席'
    if cn.startswith('达米安') or 'Damian' in en:
        return '霍桑家继承人（反派）'
    if cn.startswith('维克多') or 'Victor' in en:
        return '雷蒙德家继承人（反派）'
    if cn.startswith('亚德里安') or 'Adrian' in en:
        return '布莱克伍德家继承人（反派）'
    return '剧中主要人物'


# ---------- 生成Excel ----------
def make_excel(data, out_path, title_label=''):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '首次出场整理'

    title_font = Font(name='微软雅黑', size=16, bold=True, color='FFFFFF')
    module_font = Font(name='微软雅黑', size=13, bold=True, color='FFFFFF')
    header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
    body_font = Font(name='微软雅黑', size=11)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left = Alignment(horizontal='left', vertical='center', wrap_text=True)

    title_fill = PatternFill('solid', start_color='8B0000')
    module_fill = PatternFill('solid', start_color='C00000')
    header_fill = PatternFill('solid', start_color='E36C0A')
    alt_fill = PatternFill('solid', start_color='FDE9D9')

    thin = Side(style='thin', color='999999')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    widths = {'A': 8, 'B': 18, 'C': 20, 'D': 14, 'E': 26, 'F': 60}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    row = 1
    base_title = title_label or (data.get('title') or '剧本')
    ws.merge_cells(f'A{row}:F{row}')
    c = ws.cell(row, 1, f'《{base_title}》首次出场人物与场景整理')
    c.font = title_font
    c.fill = title_fill
    c.alignment = center
    ws.row_dimensions[row].height = 34
    row += 2

    # 人物模块
    ws.merge_cells(f'A{row}:F{row}')
    c = ws.cell(row, 1, '【人物】')
    c.font = module_font
    c.fill = module_fill
    c.alignment = center
    ws.row_dimensions[row].height = 26
    row += 1

    headers = ['序号', '人物名称', '英文名', '首次出场集数', '身份/角色定位', '人物背景分析']
    for j, h in enumerate(headers, 1):
        c = ws.cell(row, j, h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center
        c.border = border
    ws.row_dimensions[row].height = 22
    row += 1

    for i, item in enumerate(data['people'], 1):
        # 兼容4元素(中文剧本)和5元素(双语剧本, 含英文名)
        if len(item) == 5:
            name, en_name, ep, role, brief = item
        else:
            name, ep, role, brief = item
            en_name = ''
        vals = [i, name, en_name, f'第{zh_num_to_cn(ep)}集', role, brief]
        for j, v in enumerate(vals, 1):
            c = ws.cell(row, j, v)
            c.font = body_font
            c.border = border
            c.alignment = center if j in (1, 2, 3, 4, 5) else left
            if i % 2 == 0:
                c.fill = alt_fill
        ws.row_dimensions[row].height = 46
        row += 1

    row += 2

    # 场景模块
    ws.merge_cells(f'A{row}:F{row}')
    c = ws.cell(row, 1, '【场景】')
    c.font = module_font
    c.fill = module_fill
    c.alignment = center
    ws.row_dimensions[row].height = 26
    row += 1

    headers2 = ['序号', '场景名称', '英文名', '首次出场集数', '场景类型', '场景说明']
    for j, h in enumerate(headers2, 1):
        c = ws.cell(row, j, h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center
        c.border = border
    ws.row_dimensions[row].height = 22
    row += 1

    for i, (name, ep, typ) in enumerate(data['scenes'], 1):
        desc = scene_desc(name)
        vals = [i, name, '', f'第{zh_num_to_cn(ep)}集', ('内景' if typ in ('内', '室内', '内景') else '外景'), desc]
        for j, v in enumerate(vals, 1):
            c = ws.cell(row, j, v)
            c.font = body_font
            c.border = border
            c.alignment = center if j in (1, 2, 3, 4, 5) else left
            if i % 2 == 0:
                c.fill = alt_fill
        ws.row_dimensions[row].height = 26
        row += 1

    # 字幕模块（若有）
    subs = data.get('subtitles') or []
    if subs:
        row += 1
        ws.merge_cells(f'A{row}:F{row}')
        c = ws.cell(row, 1, '【字幕】')
        c.font = module_font
        c.fill = module_fill
        c.alignment = center
        ws.row_dimensions[row].height = 26
        row += 1
        headers3 = ['序号', '字幕内容', '', '出现集数', '', '说明']
        for j, h in enumerate(headers3, 1):
            c = ws.cell(row, j, h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = center
            c.border = border
        ws.row_dimensions[row].height = 22
        row += 1
        for i, (ep, sub) in enumerate(subs, 1):
            vals = [i, sub, '', f'第{zh_num_to_cn(ep)}集', '', '剧中字幕标注']
            for j, v in enumerate(vals, 1):
                c = ws.cell(row, j, v)
                c.font = body_font
                c.border = border
                c.alignment = center if j in (1, 2, 3, 4, 5) else left
                if i % 2 == 0:
                    c.fill = alt_fill
            ws.row_dimensions[row].height = 22
            row += 1

    ws.freeze_panes = 'A3'
    wb.save(out_path)
    return out_path


def make_excel_macro(data, out_path, excel_app=None):
    """生成含'点击复制'宏的Excel(.xlsm)。
    单击B列(人物名称)或C列(英文名)单元格时自动复制内容。
    依赖本机Excel + pywin32；若不可用返回 None，调用方应回退到普通xlsx。"""
    import os
    import tempfile
    out_path = os.path.abspath(out_path)
    # 1) 先用 openpyxl 生成数据到临时 xlsx
    tmp_fd, tmp_raw = tempfile.mkstemp(suffix='.xlsx')
    os.close(tmp_fd)
    os.remove(tmp_raw)
    tmp_xlsx = os.path.splitext(tmp_raw)[0] + '_data.xlsx'
    make_excel(data, tmp_xlsx)
    own_app = False
    try:
        # 2) 用 COM 打开、注入宏、另存为 xlsm（动态导入，避免PyInstaller打包win32com）
        import importlib
        win32 = importlib.import_module('win32com.client')
        if excel_app is None:
            excel_app = win32.DispatchEx('Excel.Application')
            excel_app.Visible = False
            excel_app.DisplayAlerts = False
            own_app = True
        wb = excel_app.Workbooks.Open(tmp_xlsx)
        vba = (
            'Private Sub Worksheet_SelectionChange(ByVal Target As Range)\n'
            '    On Error Resume Next\n'
            '    If Not Intersect(Target, Range("B:C")) Is Nothing Then\n'
            '        If Target.Cells.Count = 1 Then\n'
            '            If Len(Trim(Target.Value)) > 0 Then\n'
            '                Target.Copy\n'
            '            End If\n'
            '        End If\n'
            '    End If\n'
            '    On Error GoTo 0\n'
            'End Sub\n'
        )
        vbproj = wb.VBProject
        written = False
        for i in range(1, vbproj.VBComponents.Count + 1):
            comp = vbproj.VBComponents(i)
            if comp.Name.startswith('Sheet'):
                comp.CodeModule.AddFromString(vba)
                written = True
                break
        if not written:
            for i in range(1, vbproj.VBComponents.Count + 1):
                comp = vbproj.VBComponents(i)
                if comp.Type == 100:
                    comp.CodeModule.AddFromString(vba)
                    break
        if os.path.exists(out_path):
            os.remove(out_path)
        wb.SaveAs(out_path, 52)
        wb.Close(False)
        return out_path
    except Exception:
        # COM 不可用，返回 None，调用方回退到普通 xlsx
        return None
    finally:
        if own_app and 'excel_app' in dir() and excel_app is not None:
            try:
                excel_app.Quit()
            except Exception:
                pass
        if os.path.exists(tmp_xlsx):
            try:
                os.remove(tmp_xlsx)
            except Exception:
                pass
    """批量处理汇总表。summaries: [{file, people_n, scene_n, sub_n, people_list}]"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '批量汇总'

    title_font = Font(name='微软雅黑', size=14, bold=True, color='FFFFFF')
    header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
    body_font = Font(name='微软雅黑', size=10)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left = Alignment(horizontal='left', vertical='center', wrap_text=True)
    title_fill = PatternFill('solid', start_color='8B0000')
    header_fill = PatternFill('solid', start_color='E36C0A')
    alt_fill = PatternFill('solid', start_color='FDE9D9')
    thin = Side(style='thin', color='999999')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    widths = {'A': 10, 'B': 40, 'C': 10, 'D': 10, 'E': 10, 'F': 70}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    ws.merge_cells('A1:F1')
    c = ws.cell(1, 1, '剧本批量解析汇总')
    c.font = title_font
    c.fill = title_fill
    c.alignment = center
    ws.row_dimensions[1].height = 30
    ws.row_dimensions[2].height = 6

    headers = ['序号', '剧本文件', '人物数', '场景数', '字幕数', '主要角色（中文/英文）']
    for j, h in enumerate(headers, 1):
        c = ws.cell(3, j, h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center
        c.border = border
    ws.row_dimensions[3].height = 22

    row = 4
    for i, s in enumerate(summaries, 1):
        vals = [i, s['file'], s['people_n'], s['scene_n'], s['sub_n'], s['people_list']]
        for j, v in enumerate(vals, 1):
            c = ws.cell(row, j, v)
            c.font = body_font
            c.border = border
            c.alignment = center if j in (1, 3, 4, 5) else left
            if i % 2 == 0:
                c.fill = alt_fill
        ws.row_dimensions[row].height = 30
        row += 1

    ws.freeze_panes = 'A4'
    wb.save(out_path)
    return out_path


def make_merged_excel(all_data, out_path):
    """批量合并表：把多个剧本的完整数据写进同一工作簿。
    all_data: [{name, data}]，data 为 parse_script 的返回结构。
    结构：Sheet1=总览(每剧本摘要+跳转)，之后每个剧本一个sheet(人物/场景/字幕)。"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = openpyxl.Workbook()
    # 删除默认sheet，稍后重建
    wb.remove(wb.active)

    title_font = Font(name='微软雅黑', size=14, bold=True, color='FFFFFF')
    module_font = Font(name='微软雅黑', size=12, bold=True, color='FFFFFF')
    header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
    body_font = Font(name='微软雅黑', size=10)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left = Alignment(horizontal='left', vertical='center', wrap_text=True)
    title_fill = PatternFill('solid', start_color='8B0000')
    module_fill = PatternFill('solid', start_color='C00000')
    header_fill = PatternFill('solid', start_color='E36C0A')
    alt_fill = PatternFill('solid', start_color='FDE9D9')
    thin = Side(style='thin', color='999999')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ---- Sheet1: 总览 ----
    ov = wb.create_sheet('总览')
    widths = {'A': 8, 'B': 42, 'C': 10, 'D': 10, 'E': 10, 'F': 70}
    for col, w in widths.items():
        ov.column_dimensions[col].width = w
    ov.merge_cells('A1:F1')
    c = ov.cell(1, 1, '剧本批量解析总览（点击下方工作表标签查看各剧本完整数据）')
    c.font = title_font
    c.fill = title_fill
    c.alignment = center
    ov.row_dimensions[1].height = 30
    headers = ['序号', '剧本文件', '人物数', '场景数', '字幕数', '主要角色（中文/英文）']
    for j, h in enumerate(headers, 1):
        c = ov.cell(3, j, h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center
        c.border = border
    ov.row_dimensions[3].height = 22
    r = 4
    for i, item in enumerate(all_data, 1):
        d = item['data']
        plist = []
        for p in d['people'][:12]:
            if len(p) == 5:
                cn, en, ep, role, _ = p
            else:
                cn, ep, role, _ = p
                en = ''
            name = (cn or en or '')
            if en and cn:
                name = f'{cn}/{en}'
            plist.append(name)
        pstr = '、'.join(plist) + ('…' if len(d['people']) > 12 else '')
        vals = [i, item['name'], len(d['people']), len(d['scenes']), len(d.get('subtitles', [])), pstr]
        for j, v in enumerate(vals, 1):
            c = ov.cell(r, j, v)
            c.font = body_font
            c.border = border
            c.alignment = center if j in (1, 3, 4, 5) else left
            if i % 2 == 0:
                c.fill = alt_fill
        ov.row_dimensions[r].height = 28
        r += 1
    ov.freeze_panes = 'A4'

    # ---- 每个剧本一个 sheet ----
    for idx, item in enumerate(all_data, 1):
        d = item['data']
        sheet_title = f'{idx}. ' + (item.get('short') or item['name'][:20])
        ws = wb.create_sheet(sheet_title[:31])
        ws_cols = {'A': 8, 'B': 18, 'C': 20, 'D': 14, 'E': 26, 'F': 55}
        for col, w in ws_cols.items():
            ws.column_dimensions[col].width = w
        row = 1
        ws.merge_cells(f'A{row}:F{row}')
        c = ws.cell(row, 1, f'《{item["name"]}》')
        c.font = title_font
        c.fill = title_fill
        c.alignment = center
        ws.row_dimensions[row].height = 30
        row += 2

        # 人物模块
        ws.merge_cells(f'A{row}:F{row}')
        c = ws.cell(row, 1, '【人物】')
        c.font = module_font
        c.fill = module_fill
        c.alignment = center
        ws.row_dimensions[row].height = 24
        row += 1
        p_headers = ['序号', '人物名称', '英文名', '首次出场集数', '身份/角色定位', '人物背景分析']
        for j, h in enumerate(p_headers, 1):
            c = ws.cell(row, j, h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = center
            c.border = border
        ws.row_dimensions[row].height = 20
        row += 1
        for i, itemp in enumerate(d['people'], 1):
            if len(itemp) == 5:
                cn, en, ep, role, brief = itemp
            else:
                cn, ep, role, brief = itemp
                en = ''
            vals = [i, cn, en, f'第{zh_num_to_cn(ep)}集', role, brief]
            for j, v in enumerate(vals, 1):
                c = ws.cell(row, j, v)
                c.font = body_font
                c.border = border
                c.alignment = center if j in (1, 2, 3, 4, 5) else left
                if i % 2 == 0:
                    c.fill = alt_fill
            ws.row_dimensions[row].height = 34
            row += 1

        row += 1
        # 场景模块
        ws.merge_cells(f'A{row}:F{row}')
        c = ws.cell(row, 1, '【场景】')
        c.font = module_font
        c.fill = module_fill
        c.alignment = center
        ws.row_dimensions[row].height = 24
        row += 1
        s_headers = ['序号', '场景名称', '', '首次出场集数', '场景类型', '场景说明']
        for j, h in enumerate(s_headers, 1):
            c = ws.cell(row, j, h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = center
            c.border = border
        ws.row_dimensions[row].height = 20
        row += 1
        for i, (name, ep, typ) in enumerate(d['scenes'], 1):
            t = '内景' if typ in ('内', '室内', '内景') else '外景'
            vals = [i, name, '', f'第{zh_num_to_cn(ep)}集', t, scene_desc(name)]
            for j, v in enumerate(vals, 1):
                c = ws.cell(row, j, v)
                c.font = body_font
                c.border = border
                c.alignment = center if j in (1, 2, 3, 4, 5) else left
                if i % 2 == 0:
                    c.fill = alt_fill
            ws.row_dimensions[row].height = 22
            row += 1

        # 字幕模块
        subs = d.get('subtitles') or []
        if subs:
            row += 1
            ws.merge_cells(f'A{row}:F{row}')
            c = ws.cell(row, 1, '【字幕】')
            c.font = module_font
            c.fill = module_fill
            c.alignment = center
            ws.row_dimensions[row].height = 24
            row += 1
            for j, h in enumerate(['序号', '字幕内容', '', '出现集数', '', ''], 1):
                c = ws.cell(row, j, h)
                c.font = header_font
                c.fill = header_fill
                c.alignment = center
                c.border = border
            ws.row_dimensions[row].height = 20
            row += 1
            for i, (ep, sub) in enumerate(subs, 1):
                vals = [i, sub, '', f'第{zh_num_to_cn(ep)}集', '', '']
                for j, v in enumerate(vals, 1):
                    c = ws.cell(row, j, v)
                    c.font = body_font
                    c.border = border
                    c.alignment = center if j in (1, 2, 3, 4) else left
                    if i % 2 == 0:
                        c.fill = alt_fill
                ws.row_dimensions[row].height = 20
                row += 1

    wb.save(out_path)
    return out_path


def scene_desc(name):
    if '街道' in name:
        return '开篇外景，主角在街头初次登场，众人相遇。'
    if '酒店浴室' in name:
        return '内景，主角洗漱打理、显露真容，人物关系初现。'
    if '酒店套房' in name:
        return '内景，主角与同伴密谈，反派上门施压，身份初露端倪。'
    if name.endswith('车内'):
        return '外景，角色于车队中接获消息、推进剧情。'
    if '餐厅' in name:
        return '内景，高档餐厅，主角被对手当众挑衅，冲突爆发。'
    if '民政局' in name:
        return '外景，主角领证结婚，随后身份之谜揭晓。'
    if '酒店外' in name:
        return '外景/门口，宴会开场，宾客入场，主角受群嘲。'
    if '酒店' in name:
        return '内景，宴会主会场，正反派交锋的核心场景。'
    if '李家别墅' in name:
        return '内景，主角登门见家长，遭女方家人百般刁难。'
    if '别墅卧室' in name:
        return '内景，主角为女方疗伤上药，二人关系升温。'
    if '别墅院落' in name or '别墅外' in name:
        return '外景，下属前来禀报要事，主角布局剧情。'
    if '紫金天宫门外' in name:
        return '内景，国宴门口，主角持请柬到场，遭对手质疑。'
    if '紫金天宫' in name:
        return '内景，丝路国宴主会场，四国特使齐至，剧情推向高潮。'
    return '剧本场景。'


# ---------- GUI 界面 ----------
def gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk

    # 打包成 exe 时，把启动异常写入日志便于诊断
    def _log_err(msg):
        try:
            import traceback
            import tempfile
            base = getattr(sys, '_MEIPASS', None) or tempfile.gettempdir()
            logpath = os.path.join(base, 'tool_error.log')
            with open(logpath, 'a', encoding='utf-8') as f:
                f.write(msg + '\n')
                f.write(traceback.format_exc() + '\n')
        except Exception:
            pass

    is_exe = bool(getattr(sys, 'frozen', False))
    missing = ensure_deps()
    # exe 环境不应弹"缺少依赖"（库已打包）；若确实缺则记录日志后仍继续
    if missing and not is_exe:
        # 用 withdraw 的临时窗口弹窗（确保销毁，不留多余Tk窗口）
        tmp = tk.Tk()
        tmp.withdraw()
        ret = messagebox.askyesno(
            '缺少依赖',
            f'缺少运行库: {", ".join(missing)}\n\n是否自动安装（需联网）？\n否则请手动安装后再运行。')
        tmp.destroy()
        if ret:
            for m in missing:
                try:
                    auto_install(m)
                except Exception as e:
                    messagebox.showerror('安装失败', f'自动安装 {m} 失败:\n{e}\n请手动执行: pip install {m}')
                    return

    # 启用拖拽支持：直接用 TkinterDnD.Tk() 创建主窗口（仅一个 Tk 实例）
    DND_OK = False
    root = None
    try:
        from tkinterdnd2 import TkinterDnD
        root = TkinterDnD.Tk()
        # 验证 tkdnd 加载成功；失败则销毁此窗口，改用普通 Tk
        root.tk.eval('package require tkdnd')
        DND_OK = True
        _log_err('[init] tkdnd加载OK')
    except Exception:
        # tkdnd 不可用：销毁已创建的窗口，确保只保留一个 Tk
        if root is not None:
            try:
                root.destroy()
            except Exception:
                pass
            root = None
        root = tk.Tk()
        DND_OK = False
    root.title('剧本首次出场人物与场景整理工具')
    root.configure(bg='#f5f5f5')

    # 让窗口自适应内容高度并居中，避免按钮被挤出屏幕
    def center_window(win, w, h=None):
        win.update_idletasks()
        sw = win.winfo_screenwidth()
        sh = win.winfo_screenheight()
        if h is None:
            req = win.winfo_reqheight()
            h = min(req + 30, sh - 60)
        x = max((sw - w) // 2, 0)
        y = max((sh - h) // 2, 0)
        win.geometry(f'{w}x{h}+{x}+{y}')

    try:
        # 应用软件图标
        ico = _resource_path('icon.ico')
        if os.path.exists(ico):
            root.iconbitmap(ico)
        else:
            root.iconbitmap(default='')
    except Exception:
        pass

    # 顶部标题（品牌配色 + 副标题）
    header = tk.Frame(root, bg='#7a1f1f')
    header.pack(fill='x')
    tk.Label(header, text='🎬 剧本解析工具', font=('微软雅黑', 16, 'bold'),
             bg='#7a1f1f', fg='white').pack(pady=(8, 0))
    tk.Label(header, text='自动识别人物 · 场景 · 时间字幕，导出剪辑资料表',
             font=('微软雅黑', 9), bg='#7a1f1f', fg='#f0d0c0').pack(pady=(2, 8))

    # 主内容区
    main = tk.Frame(root, bg='#f5f5f5')
    main.pack(fill='x', padx=14, pady=6)

    # === 第一组：文件选择 ===
    group1 = tk.LabelFrame(main, text=' ① 选择剧本文件 ', font=('微软雅黑', 10, 'bold'),
                           bg='#f5f5f5', fg='#7a1f1f', bd=1, relief='groove', padx=10, pady=8)
    group1.pack(fill='x', pady=(0, 8))

    frm = tk.Frame(group1, bg='#f5f5f5')
    frm.pack(fill='x')

    def choose_files():
        fps = filedialog.askopenfilenames(filetypes=[('Word文档', '*.docx'), ('所有文件', '*.*')])
        for fp in fps:
            if fp not in files_list:
                files_list.append(fp)
                listbox.insert('end', os.path.basename(fp))
        update_count()

    tk.Button(frm, text='＋ 添加文件（可多选）', command=choose_files, font=('微软雅黑', 10),
              bg='#E36C0A', fg='white', activebackground='#c95a00',
              padx=14, pady=3, cursor='hand2', relief='flat').pack(side='left')
    tk.Label(frm, text='  或直接拖拽 docx 文件到下方列表', font=('微软雅黑', 9), fg='#999',
             bg='#f5f5f5').pack(side='left', padx=8)

    # 文件列表区
    listframe = tk.Frame(group1, bg='#f5f5f5')
    listframe.pack(fill='x', pady=6)
    files_list = []
    listbox = tk.Listbox(listframe, height=3, font=('微软雅黑', 9), selectmode='extended',
                         bg='#ffffff', relief='solid', bd=1, highlightthickness=1,
                         highlightcolor='#E36C0A', highlightbackground='#dddddd')
    listbox.pack(side='left', fill='x', expand=True, padx=(0, 4))
    sb = tk.Scrollbar(listframe, orient='vertical', command=listbox.yview)
    listbox.configure(yscrollcommand=sb.set)
    sb.pack(side='right', fill='y')

    list_btns = tk.Frame(group1, bg='#f5f5f5')
    list_btns.pack(fill='x')
    def remove_selected():
        sel = listbox.curselection()
        for idx in reversed(sel):
            listbox.delete(idx)
            del files_list[idx]
        update_count()
    def clear_list():
        listbox.delete(0, 'end')
        files_list.clear()
        update_count()
    tk.Button(list_btns, text='移除选中', command=remove_selected, font=('微软雅黑', 9),
              bg='#efefef', relief='flat', padx=10).pack(side='left')
    tk.Button(list_btns, text='清空列表', command=clear_list, font=('微软雅黑', 9),
              bg='#efefef', relief='flat', padx=10).pack(side='left', padx=6)
    count_var = tk.StringVar(value='已选 0 个文件')
    tk.Label(list_btns, textvariable=count_var, font=('微软雅黑', 9), fg='#666',
             bg='#f5f5f5').pack(side='right')
    def update_count():
        count_var.set(f'已选 {len(files_list)} 个文件')

    # === 第二组：解析设置 ===
    group2 = tk.LabelFrame(main, text=' ② 解析设置 ', font=('微软雅黑', 10, 'bold'),
                           bg='#f5f5f5', fg='#7a1f1f', bd=1, relief='groove', padx=10, pady=6)
    group2.pack(fill='x', pady=(0, 8))

    # 集数 + 语言 一行
    row1 = tk.Frame(group2, bg='#f5f5f5')
    row1.pack(fill='x', pady=2)
    tk.Label(row1, text='提取前多少集:', font=('微软雅黑', 10), bg='#f5f5f5').pack(side='left')
    ep_var = tk.StringVar(value='')
    tk.Entry(row1, textvariable=ep_var, width=6, font=('微软雅黑', 10)).pack(side='left', padx=4)
    tk.Label(row1, text='(留空=全本)  ', font=('微软雅黑', 9), fg='#999', bg='#f5f5f5').pack(side='left')
    tk.Label(row1, text='角色名语言:', font=('微软雅黑', 10), bg='#f5f5f5').pack(side='left', padx=(10, 0))
    lang_var = tk.StringVar(value='自动')
    ttk.Combobox(row1, textvariable=lang_var, width=6, state='readonly',
                 values=('自动', '英文', '中文'), font=('微软雅黑', 10)).pack(side='left', padx=4)

    # 宏选项
    row2 = tk.Frame(group2, bg='#f5f5f5')
    row2.pack(fill='x', pady=2)
    macro_var = tk.BooleanVar(value=True)
    tk.Checkbutton(row2, text='生成"点击角色名即可复制"的Excel（.xlsm）', variable=macro_var,
                   font=('微软雅黑', 9), bg='#f5f5f5', activebackground='#f5f5f5',
                   selectcolor='#ffffff').pack(side='left')
    tk.Label(row2, text='（需本机Excel，首次打开需启用宏）', font=('微软雅黑', 8), fg='#999',
             bg='#f5f5f5').pack(side='left', padx=6)

    # 校准按钮
    row3 = tk.Frame(group2, bg='#f5f5f5')
    row3.pack(fill='x', pady=2)
    tk.Button(row3, text='⚙ 校准规则…', command=lambda: open_calibration(),
              font=('微软雅黑', 9), bg='#8B0000', fg='white', activebackground='#6b0000',
              relief='flat', padx=12, pady=2, cursor='hand2').pack(side='left')
    tk.Label(row3, text='（排除角色 / 名字对应 / 身份覆盖，解析时自动应用）', font=('微软雅黑', 8),
             fg='#999', bg='#f5f5f5').pack(side='left', padx=8)

    # 拖拽文件添加
    def drop_files(event):
        import re as _re
        data = event.data
        # 处理 tkinterdnd2 返回的文件路径列表
        paths = []
        if data:
            # 支持花括号包裹的路径和空格分隔
            parts = _re.findall(r'\{[^}]+\}|[^\s]+', data)
            for p in parts:
                p = p.strip('{}').strip()
                if p.lower().endswith('.docx'):
                    paths.append(p)
        for fp in paths:
            if os.path.exists(fp) and fp not in files_list:
                files_list.append(fp)
                listbox.insert('end', os.path.basename(fp))
        update_count()

    def drop_files_clean(event):
        return 'break'

    if DND_OK:
        try:
            root.drop_target_register('DND_Files')
            root.dnd_bind('<<Drop>>', drop_files)
            listbox.drop_target_register('DND_Files')
            listbox.dnd_bind('<<Drop>>', drop_files)
        except Exception:
            pass

    # 日志区（放入卡片分组）
    group3 = tk.LabelFrame(main, text=' ③ 处理日志 ', font=('微软雅黑', 10, 'bold'),
                           bg='#f5f5f5', fg='#7a1f1f', bd=1, relief='groove', padx=10, pady=6)
    group3.pack(fill='x')
    logframe = tk.Frame(group3, bg='#f5f5f5')
    logframe.pack(fill='x')
    log = tk.Text(logframe, height=4, font=('微软雅黑', 9), bg='#ffffff', wrap='char',
                  relief='solid', bd=1, highlightthickness=1, highlightcolor='#E36C0A',
                  highlightbackground='#dddddd')
    log.pack(side='left', fill='both', expand=True, padx=(0, 4))
    log_sb = tk.Scrollbar(logframe, orient='vertical', command=log.yview)
    log.configure(yscrollcommand=log_sb.set)
    log_sb.pack(side='right', fill='y')
    log.insert('1.0', '就绪。请选择剧本文件后点击"开始整理"。\n')
    log.config(state='disabled')

    def log_msg(s):
        log.config(state='normal')
        log.insert('end', s + '\n')
        log.see('end')
        log.config(state='disabled')
        root.update()

    def open_calibration():
        """校准规则编辑窗口：排除角色 / 名字对应 / 身份覆盖"""
        cfg = load_calibration()
        win = tk.Toplevel(root)
        win.title('校准规则')
        win.configure(bg='#f5f5f5')
        win.transient(root)
        win.grab_set()
        w = min(720, win.winfo_screenwidth() - 60)
        h = min(600, win.winfo_screenheight() - 80)
        win.geometry(f'{w}x{h}+{max((win.winfo_screenwidth()-w)//2,0)}+{max((win.winfo_screenheight()-h)//2,0)}')

        tk.Label(win, text='校准规则（解析时自动应用，保存到 calibration.json）',
                 font=('微软雅黑', 12, 'bold'), bg='#8B0000', fg='white').pack(fill='x', ipady=8)

        nb = ttk.Notebook(win)
        nb.pack(fill='both', expand=True, padx=10, pady=8)

        # ---- 排除角色页 ----
        pe = tk.Frame(nb, bg='white')
        nb.add(pe, text=f'排除角色（{len(cfg["exclude"])}）')
        pe_lb = tk.Listbox(pe, height=10, font=('微软雅黑', 10))
        pe_lb.pack(fill='both', expand=True, padx=6, pady=6)
        for x in cfg['exclude']:
            pe_lb.insert('end', x)
        pe_btns = tk.Frame(pe, bg='white')
        pe_btns.pack(pady=6)
        pe_entry = tk.Entry(pe_btns, width=25, font=('微软雅黑', 10))
        pe_entry.pack(side='left', padx=4)
        def pe_add():
            v = pe_entry.get().strip()
            if v and v not in cfg['exclude']:
                cfg['exclude'].append(v)
                pe_lb.insert('end', v)
                pe_entry.delete(0, 'end')
        def pe_del():
            sel = pe_lb.curselection()
            if sel:
                idx = sel[0]
                cfg['exclude'].pop(idx)
                pe_lb.delete(idx)
        tk.Button(pe_btns, text='添加', command=pe_add, font=('微软雅黑', 9)).pack(side='left', padx=2)
        tk.Button(pe_btns, text='删除', command=pe_del, font=('微软雅黑', 9)).pack(side='left', padx=2)
        tk.Label(pe, text='这些名字不会出现在结果中（如误识别为角色的群演）', font=('微软雅黑', 9), fg='#888', bg='white').pack(pady=2)

        # ---- 名字对应页 ----
        pa = tk.Frame(nb, bg='white')
        nb.add(pa, text=f'名字对应（{len(cfg["aliases"])}）')
        pa_lb = tk.Listbox(pa, height=10, font=('微软雅黑', 10))
        pa_lb.pack(fill='both', expand=True, padx=6, pady=6)
        for old, new in cfg['aliases'].items():
            pa_lb.insert('end', f'{old} → {new}')
        pa_btns = tk.Frame(pa, bg='white')
        pa_btns.pack(pady=6)
        tk.Label(pa_btns, text='原名字:', font=('微软雅黑', 9), bg='white').pack(side='left')
        pa_old = tk.Entry(pa_btns, width=12, font=('微软雅黑', 10))
        pa_old.pack(side='left', padx=2)
        tk.Label(pa_btns, text='→ 改成:', font=('微软雅黑', 9), bg='white').pack(side='left')
        pa_new = tk.Entry(pa_btns, width=12, font=('微软雅黑', 10))
        pa_new.pack(side='left', padx=2)
        def pa_add():
            o = pa_old.get().strip()
            n = pa_new.get().strip()
            if o and n:
                cfg['aliases'][o] = n
                pa_lb.insert('end', f'{o} → {n}')
                pa_old.delete(0, 'end'); pa_new.delete(0, 'end')
        def pa_del():
            sel = pa_lb.curselection()
            if sel:
                idx = sel[0]
                line = pa_lb.get(idx)
                key = line.split(' → ')[0].strip()
                if key in cfg['aliases']:
                    del cfg['aliases'][key]
                pa_lb.delete(idx)
        tk.Button(pa_btns, text='添加', command=pa_add, font=('微软雅黑', 9)).pack(side='left', padx=2)
        tk.Button(pa_btns, text='删除', command=pa_del, font=('微软雅黑', 9)).pack(side='left', padx=2)
        tk.Label(pa, text='把识别错误/不统一的角色名统一（如 George → 乔治）', font=('微软雅黑', 9), fg='#888', bg='white').pack(pady=2)

        # ---- 身份覆盖页 ----
        pr = tk.Frame(nb, bg='white')
        nb.add(pr, text=f'身份覆盖（{len(cfg["roles"])}）')
        pr_lb = tk.Listbox(pr, height=10, font=('微软雅黑', 10))
        pr_lb.pack(fill='both', expand=True, padx=6, pady=6)
        for nm, role in cfg['roles'].items():
            pr_lb.insert('end', f'{nm} → {role}')
        pr_btns = tk.Frame(pr, bg='white')
        pr_btns.pack(pady=6)
        tk.Label(pr_btns, text='角色名:', font=('微软雅黑', 9), bg='white').pack(side='left')
        pr_nm = tk.Entry(pr_btns, width=12, font=('微软雅黑', 10))
        pr_nm.pack(side='left', padx=2)
        tk.Label(pr_btns, text='→ 身份:', font=('微软雅黑', 9), bg='white').pack(side='left')
        pr_role = tk.Entry(pr_btns, width=15, font=('微软雅黑', 10))
        pr_role.pack(side='left', padx=2)
        def pr_add():
            nm = pr_nm.get().strip()
            rl = pr_role.get().strip()
            if nm and rl:
                cfg['roles'][nm] = rl
                pr_lb.insert('end', f'{nm} → {rl}')
                pr_nm.delete(0, 'end'); pr_role.delete(0, 'end')
        def pr_del():
            sel = pr_lb.curselection()
            if sel:
                idx = sel[0]
                line = pr_lb.get(idx)
                key = line.split(' → ')[0].strip()
                if key in cfg['roles']:
                    del cfg['roles'][key]
                pr_lb.delete(idx)
        tk.Button(pr_btns, text='添加', command=pr_add, font=('微软雅黑', 9)).pack(side='left', padx=2)
        tk.Button(pr_btns, text='删除', command=pr_del, font=('微软雅黑', 9)).pack(side='left', padx=2)
        tk.Label(pr, text='给特定角色指定身份定位（如 山本康盛 → 头号反派）', font=('微软雅黑', 9), fg='#888', bg='white').pack(pady=2)

        # 底部按钮
        btnf = tk.Frame(win, bg='#f5f5f5')
        btnf.pack(pady=8)
        def save():
            if save_calibration(cfg):
                messagebox.showinfo('保存', '校准规则已保存，下次解析自动应用。')
                win.destroy()
            else:
                messagebox.showerror('保存', '保存失败，请检查目录权限。')
        tk.Button(btnf, text='保存', command=save, font=('微软雅黑', 11, 'bold'),
                  bg='#8B0000', fg='white', padx=30).pack(side='left', padx=10)
        tk.Button(btnf, text='取消', command=win.destroy, font=('微软雅黑', 11),
                  bg='#cccccc', padx=30).pack(side='left', padx=10)

    def process_one_file(fp, max_ep, is_full, lang_mode, use_macro):
        """解析单个文件并导出Excel，返回 {file, path, data} 或 None(失败)"""
        log_msg(f'  解析: {os.path.basename(fp)}')
        try:
            data, err = parse_script(fp, max_ep, lang_mode)
            if data and not err:
                apply_calibration(data)
        except Exception as e:
            log_msg('    出错: ' + str(e))
            return None
        if err:
            log_msg('    失败: ' + err)
            return None
        if not data['people'] and not data['scenes']:
            log_msg('    未提取到人物或场景')
            return None
        suffix = '全本' if is_full else f'前{max_ep}集'
        # 导出：若用宏则生成 .xlsm，否则普通 .xlsx
        if use_macro:
            out = os.path.splitext(fp)[0] + f'_首次出场_{suffix}.xlsm'
            result = make_excel_macro(data, out)
            if not result:
                # 宏不可用，回退 xlsx
                log_msg('    本机Excel不可用，已改用普通xlsx')
                out = os.path.splitext(fp)[0] + f'_首次出场_{suffix}.xlsx'
                make_excel(data, out)
        else:
            out = os.path.splitext(fp)[0] + f'_首次出场_{suffix}.xlsx'
            make_excel(data, out)
        log_msg(f'    完成：人物{len(data["people"])} 场景{len(data["scenes"])} 字幕{len(data.get("subtitles", []))}')
        log_msg(f'    已生成: {os.path.basename(out)}')
        # 返回完整 data 供合并表使用
        return {
            'file': os.path.basename(fp),
            'path': fp,
            'data': data,
        }

    def run():
        if not files_list:
            messagebox.showwarning('提示', '请先选择至少一个剧本文件')
            return
        use_macro = macro_var.get()
        # 集数：留空=全本(用极大值), 否则取数字
        ep_input = ep_var.get().strip()
        is_full = False
        if ep_input:
            try:
                max_ep = int(ep_input)
            except ValueError:
                messagebox.showerror('错误', '集数必须是数字，或留空表示全本')
                return
            if max_ep <= 0:
                is_full = True
                max_ep = 10**9
        else:
            is_full = True
            max_ep = 10**9
        # 语言模式
        lang_map = {'英文': 'en', '中文': 'zh', '自动': 'auto'}
        lang_mode = lang_map.get(lang_var.get(), 'auto')
        scope = '全本' if is_full else f'前{max_ep}集'
        log_msg(f'开始处理 {len(files_list)} 个剧本（{scope}）｜角色语言：{lang_var.get()}')

        if len(files_list) == 1:
            # 单文件：解析 -> 预览 -> 确认后导出
            fp = files_list[0]
            if not os.path.exists(fp):
                messagebox.showerror('错误', '文件不存在:\n' + fp)
                return
            try:
                data, err = parse_script(fp, max_ep, lang_mode)
                if data and not err:
                    apply_calibration(data)
            except Exception as e:
                log_msg('解析出错: ' + str(e))
                messagebox.showerror('解析失败', str(e))
                return
            if err:
                messagebox.showerror('解析失败', err)
                return
            if not data['people'] and not data['scenes']:
                messagebox.showwarning('提示', '未能提取到任何人物或场景，请检查剧本格式。')
                return
            log_msg(f'解析完成：人物 {len(data["people"])} 个，场景 {len(data["scenes"])} 个。')
            ok = preview_window(data, is_full, max_ep)
            if not ok:
                log_msg('已取消导出。')
                return
            suffix = '全本' if is_full else f'前{max_ep}集'
            if use_macro:
                out = os.path.splitext(fp)[0] + f'_首次出场_{suffix}.xlsm'
                result = make_excel_macro(data, out)
                if not result:
                    log_msg('本机Excel不可用，已改用普通xlsx')
                    out = os.path.splitext(fp)[0] + f'_首次出场_{suffix}.xlsx'
                    make_excel(data, out)
            else:
                out = os.path.splitext(fp)[0] + f'_首次出场_{suffix}.xlsx'
                make_excel(data, out)
            log_msg(f'完成！人物 {len(data["people"])} 个，场景 {len(data["scenes"])} 个。')
            log_msg('已生成: ' + out)
            ret = messagebox.askyesno('完成', f'已生成Excel（人物{len(data["people"])}个，场景{len(data["scenes"])}个）。\n\n是否立即打开文件？')
            if ret:
                try:
                    os.startfile(out)
                except Exception:
                    pass
        else:
            # 批量处理：逐个解析导出，最后生成汇总表
            summaries = []
            fail = 0
            for idx, fp in enumerate(files_list, 1):
                log_msg(f'[{idx}/{len(files_list)}]')
                if not os.path.exists(fp):
                    log_msg(f'    文件不存在，跳过')
                    fail += 1
                    continue
                s = process_one_file(fp, max_ep, is_full, lang_mode, use_macro)
                if s:
                    summaries.append(s)
                else:
                    fail += 1
            log_msg(f'批量处理完成：成功 {len(summaries)}，失败 {fail}。')
            # 生成合并表（每个剧本完整数据，依次排列）
            if summaries:
                default_dir = os.path.dirname(files_list[0]) or '.'
                sum_out = os.path.join(default_dir, '剧本批量解析合并.xlsx')
                try:
                    all_data = []
                    for s in summaries:
                        base = os.path.splitext(s['file'])[0]
                        all_data.append({
                            'name': s['file'],
                            'short': base,
                            'data': s['data'],
                        })
                    make_merged_excel(all_data, sum_out)
                    log_msg('合并表: ' + sum_out)
                except Exception as e:
                    log_msg('生成合并表出错: ' + str(e))
                ret = messagebox.askyesno('完成', f'批量处理完成：成功{len(summaries)}个，失败{fail}个。\n\n合并表已生成（含每个剧本的完整人物/场景/字幕，按剧本依次排列）。\n\n是否打开？')
                if ret:
                    try:
                        os.startfile(sum_out)
                    except Exception:
                        pass
            else:
                messagebox.showwarning('提示', '全部文件处理失败，请检查剧本格式。')

    def preview_window(data, is_full, max_ep):
        """预览窗口：展示提取的人物与场景，用户确认后导出，取消则不导出。返回 True=确认"""
        import tkinter as tk
        from tkinter import ttk, messagebox
        win = tk.Toplevel(root)
        win.title('预览确认')
        win.configure(bg='#f5f5f5')
        win.transient(root)
        win.grab_set()
        # 自适应并居中，避免按钮被挤出屏幕
        win.update_idletasks()
        sw = win.winfo_screenwidth()
        sh = win.winfo_screenheight()
        w = min(820, sw - 40)
        h = min(560, sh - 80)
        x = max((sw - w) // 2, 0)
        y = max((sh - h) // 2, 0)
        win.geometry(f'{w}x{h}+{x}+{y}')

        scope = '全本' if is_full else f'前{max_ep}集'
        tk.Label(win, text=f'提取结果预览（{scope}）—— 请确认后导出',
                 font=('微软雅黑', 13, 'bold'), bg='#8B0000', fg='white').pack(fill='x', ipady=8)

        nb = ttk.Notebook(win)
        nb.pack(fill='both', expand=True, padx=10, pady=8)

        # 人物页
        pframe = tk.Frame(nb, bg='white')
        nb.add(pframe, text=f'人物（{len(data["people"])}）')
        cols = ('序号', '人物名称', '英文名', '首次出场集数', '身份/角色定位')
        tree_p = ttk.Treeview(pframe, columns=cols, show='headings', height=18)
        widths = {'序号': 50, '人物名称': 140, '英文名': 150, '首次出场集数': 90, '身份/角色定位': 200}
        for c in cols:
            tree_p.heading(c, text=c)
            tree_p.column(c, width=widths[c], anchor='center')
        vs = ttk.Scrollbar(pframe, orient='vertical', command=tree_p.yview)
        tree_p.configure(yscrollcommand=vs.set)
        tree_p.pack(side='left', fill='both', expand=True)
        vs.pack(side='right', fill='y')
        for i, item in enumerate(data['people'], 1):
            if len(item) == 5:
                cn, en, ep, role, _ = item
            else:
                cn, ep, role, _ = item
                en = ''
            tree_p.insert('', 'end', values=(i, cn, en, f'第{zh_num_to_cn(ep)}集', role))

        # 场景页
        sframe = tk.Frame(nb, bg='white')
        nb.add(sframe, text=f'场景（{len(data["scenes"])}）')
        cols2 = ('序号', '场景名称', '首次出场集数', '场景类型')
        tree_s = ttk.Treeview(sframe, columns=cols2, show='headings', height=18)
        widths2 = {'序号': 50, '场景名称': 260, '首次出场集数': 100, '场景类型': 80}
        for c in cols2:
            tree_s.heading(c, text=c)
            tree_s.column(c, width=widths2[c], anchor='center')
        vs2 = ttk.Scrollbar(sframe, orient='vertical', command=tree_s.yview)
        tree_s.configure(yscrollcommand=vs2.set)
        tree_s.pack(side='left', fill='both', expand=True)
        vs2.pack(side='right', fill='y')
        for i, (name, ep, typ) in enumerate(data['scenes'], 1):
            t = '内景' if typ in ('内', '室内', '内景') else '外景'
            tree_s.insert('', 'end', values=(i, name, f'第{zh_num_to_cn(ep)}集', t))

        # 字幕页（若有）
        subs = data.get('subtitles') or []
        if subs:
            subframe = tk.Frame(nb, bg='white')
            nb.add(subframe, text=f'字幕（{len(subs)}）')
            cols3 = ('序号', '字幕内容', '出现集数')
            tree_sub = ttk.Treeview(subframe, columns=cols3, show='headings', height=18)
            widths3 = {'序号': 50, '字幕内容': 320, '出现集数': 100}
            for c in cols3:
                tree_sub.heading(c, text=c)
                tree_sub.column(c, width=widths3[c], anchor='center')
            vs3 = ttk.Scrollbar(subframe, orient='vertical', command=tree_sub.yview)
            tree_sub.configure(yscrollcommand=vs3.set)
            tree_sub.pack(side='left', fill='both', expand=True)
            vs3.pack(side='right', fill='y')
            for i, (ep, sub) in enumerate(subs, 1):
                tree_sub.insert('', 'end', values=(i, sub, f'第{zh_num_to_cn(ep)}集'))

        # 底部按钮
        btnf = tk.Frame(win, bg='#f5f5f5')
        btnf.pack(pady=10)
        result = {'ok': False}
        def confirm():
            result['ok'] = True
            win.destroy()
        def cancel():
            result['ok'] = False
            win.destroy()
        tk.Button(btnf, text='确认导出', command=confirm, font=('微软雅黑', 12, 'bold'),
                  bg='#8B0000', fg='white', padx=30).pack(side='left', padx=15)
        tk.Button(btnf, text='取消', command=cancel, font=('微软雅黑', 12),
                  bg='#cccccc', padx=30).pack(side='left', padx=15)

        # 等待窗口关闭
        win.wait_window()
        return result['ok']

    # 开始整理按钮（居中，醒目）
    btn_bar = tk.Frame(root, bg='#f5f5f5')
    btn_bar.pack(pady=(6, 10))
    tk.Button(btn_bar, text='🚀  开始整理', command=run, font=('微软雅黑', 14, 'bold'),
              bg='#E36C0A', fg='white', activebackground='#c95a00', relief='flat',
              padx=50, pady=8, cursor='hand2').pack()

    # 自适应窗口大小并居中（确保按钮可见）
    center_window(root, 640)
    try:
        root.mainloop()
    except Exception:
        _log_err('GUI 主循环异常')
        try:
            root.destroy()
        except Exception:
            pass


# ---------- 命令行入口 ----------
def cli():
    missing = ensure_deps()
    if missing:
        print('缺少依赖:', ', '.join(missing))
        if getattr(sys, 'frozen', False):
            # exe 环境不能自动安装，提示用户
            print('打包版缺少运行库，请重新打包或使用源码版。')
            return
        for m in missing:
            print(f'正在安装 {m} ...')
            auto_install(m)
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        print('用法: python script_parser.py <剧本文档.docx> [集数]')
        return
    max_ep = int(sys.argv[2]) if len(sys.argv) > 2 else 30

    data, err = parse_script(docx_path, max_ep)
    if err:
        print('解析失败:', err)
        return
    out = os.path.splitext(docx_path)[0] + f'_首次出场前{max_ep}集.xlsx'
    make_excel(data, out)
    print(f'完成! 人物{len(data["people"])}个, 场景{len(data["scenes"])}个')
    print('输出:', out)


if __name__ == '__main__':
    if len(sys.argv) > 1:
        cli()
    else:
        gui()
